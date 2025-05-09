from moviepy.editor import VideoFileClip
import whisper
from dotenv import load_dotenv
from langchain_community.document_loaders import TextLoader
from typing import List
from langchain.schema import Document
from langchain_google_genai import ChatGoogleGenerativeAI
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from fpdf import FPDF
from flask import Flask, request, jsonify, send_file,render_template,flash

app = Flask(__name__)

load_dotenv()

llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash")


def load_text_file(file_path: str) -> List[Document]:
    loader = TextLoader(file_path,encoding='utf-8')
    return loader.load()

def load_document(file_path: str) -> List[Document]:
    print('loading transcript')
    if file_path.endswith('.txt'):
        return load_text_file(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_path}")
    
def genarate_action_points(text_docs):
    # Extract the text content from the document
    print('generating action points')
    response=llm.invoke(f'This is the transcript from the meeting {text_docs[0].page_content}. Can you generate the action points from this text.And keep focus on every small detail in the transcript.')
    action_points=(response.content)
    return action_points

def save_action_points_word(action_points,name):
    print("Saving action points...")
    doc = Document()
    
    # Add a title
    title = doc.add_heading('Project Action Points', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Split the text into lines
    lines = action_points.strip().split('\n')
    
    current_level = 0
    
    for line in lines:
        # Skip empty lines
        if not line.strip():
            continue
        
        # Determine the line type and format accordingly
        line = line.strip()
        
        # Main section headers (I., II., etc.)
        if re.match(r'^[IVX]+\.', line):
            doc.add_heading(line, level=1)
            current_level = 1
            
        # Main bullet points (start with *)
        elif line.startswith('*'):
            # Check if it's a main point or sub-point based on indentation
            indent_level = len(line) - len(line.lstrip())
            
            if indent_level <= 2:  # Main bullet point
                text = line.lstrip('* ').strip()
                p = doc.add_paragraph(style='List Bullet')
                
                # Check if the text has bold formatting (surrounded by **)
                if '**' in text:
                    # Handle bold text portions
                    parts = text.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 1:  # Odd indices are bold
                            p.add_run(part).bold = True
                        else:
                            p.add_run(part)
                else:
                    p.add_run(text)
                
                current_level = 2
                
            else:  # Sub-bullet point
                text = line.lstrip('* ').strip()
                p = doc.add_paragraph(style='List Bullet 2')
                
                # Check if the text has bold formatting
                if '**' in text:
                    parts = text.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 1:  # Odd indices are bold
                            p.add_run(part).bold = True
                        else:
                            p.add_run(part)
                else:
                    p.add_run(text)
                
                current_level = 3
                
        # Regular text - treat as paragraph
        else:
            p = doc.add_paragraph(line)
    
    # Save the document
   
    doc.save(f'{name}_action_points.docx')
    print(f"Action points saved to {name}.docx")

def save_action_points_pdf(action_points,name):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=8)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.multi_cell(0, 8, action_points)  # Add action points
    pdf.output(f'{name}_action_points.pdf')
    print(f"PDF saved as {name}")


def extract_audio_from_video(video_path, audio_output_path):
    video_clip = VideoFileClip(video_path)
    audio_clip = video_clip.audio
    audio_clip.write_audiofile(audio_output_path)

def transcribe_audio(audio_path, model_size="base"):
    print("Transcribing audio...")
       
    model = whisper.load_model(model_size)
        
        # Transcribe the audio
    result = model.transcribe(audio_path)
    return result['text']    
        
def save_transcript(transcript, output_file):
    """Save the transcript in a readable format"""
    print("Saving transcript...")
    with open(output_file, 'w', encoding='utf-8') as f:
        # Write plain text format
        f.write(transcript)

@app.route('/video_to_text', methods=['POST', 'GET'])
def video_to_text(output_text_path=None, model_size="base"):
    global name
    # video_path="video1475546927.mp4"
    if request.method == 'POST':
        video_path= request.files['video']
        file_name=video_path.filename
        intermediate_audio_path = "extracted_audio.wav"
        print(file_name)
        name=file_name[:-4]
        output_text_path=f"{file_name[:-4]}_transcription.txt"
        if file_name.endswith('.mp4'):
            extract_audio_from_video(file_name, intermediate_audio_path)
        else:
            flash("Invalid file format. Please upload a video file")


        result=transcribe_audio(intermediate_audio_path, model_size)
        save_transcript(result, output_text_path)
        if os.path.exists(intermediate_audio_path):
            os.remove(intermediate_audio_path)

        text_docs = load_document(output_text_path)
        action_points=genarate_action_points(text_docs)
        save_action_points_word(action_points,name)
        save_action_points_pdf(action_points,name)
        data={"status":'success','action_points':action_points}


        return jsonify(data)
    return render_template('index.html')


@app.route('/download_word', methods=['GET'])
def download_word():
    file_path = f"{name}_action_points.docx"  # Ensure this path is correct
    return send_file(file_path, 
                     as_attachment=True, 
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/download_pdf')
def download_pdf():
        # Path to your PDF file
    pdf_path = f'{name}_action_points.pdf'
        
        # Check if file exists
    if not os.path.exists(pdf_path):
        return jsonify({'error': 'File not found'}), 404
        
        # Send the file
    return send_file(
            pdf_path, 
            mimetype='application/pdf', 
            as_attachment=True, 
            
        )


if __name__ == "__main__":
    app.run(debug=True)